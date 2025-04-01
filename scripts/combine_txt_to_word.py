"""
A script, which takes a folder of .txt files and combines them into one .docx
file. Each .txt file will start on a new page in the .docx file, titled with
the filename of the .txt file.

Requirements:
python-docx==1.1.2
"""

import os
from docx import Document
from docx.enum.section import WD_ORIENTATION


def create_word_document(source_directory: str, output_file: str):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT

    for filename in os.listdir(source_directory):
        if not filename.endswith(".txt"):
            continue

        document.add_page_break()
        document.add_heading(filename, level=1)

        with open(
            os.path.join(source_directory, filename), "r", encoding="utf-8"
        ) as file:
            content = file.read()
            document.add_paragraph(content, style="Normal")

    document.save(output_file)


directory = input("Path to the folder containing .txt files: ").strip()
output_file = input("Name of the output .docx file: ").strip()

create_word_document(directory, output_file)

print(f"Combined {len(os.listdir(directory))} .txt files into {output_file}")
